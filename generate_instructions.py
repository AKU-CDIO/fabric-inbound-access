"""Generate instruction sheet for researchers."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def generate():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\nUZIMA Fabric Data Access\n")
    run.font.size = Pt(22)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Quick Start Guide\n")
    run.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "\nEverything is pre-configured on this VM.\n"
        "Just open RStudio or Python and run the code below.\n"
    )
    run.font.size = Pt(11)

    # --- RStudio ---
    doc.add_heading("RStudio", level=2)
    doc.add_paragraph(
        "Create a new R Markdown (.Rmd) file and paste this in any code chunk:"
    )
    _code(doc, [
        "```{r}",
        "library(fabriconnect)",
        "conn <- connect_to_fabric()",
        "",
        "# List all tables in the primary lakehouse",
        "tables <- list_tables(conn)",
        "",
        "# Read a table",
        'df <- read_table(conn, "registeredparticipants")',
        "head(df)",
        "```",
    ])
    doc.add_paragraph("That's it. No login, no token setup, nothing else needed.")

    # --- Python ---
    doc.add_heading("Python", level=2)
    doc.add_paragraph("In a terminal or notebook:")
    _code(doc, [
        "from fabricpy import FabricLakehouse",
        "lh = FabricLakehouse()",
        "tables = lh.list_tables()",
        'df = lh.read_table("agents").to_pandas()',
        "df.head()",
    ])

    # --- Test ---
    doc.add_heading("Test It", level=2)
    _code(doc, [
        "# RStudio:",
        'source("Runbooks/test-r-access.R")',
        "",
        "# Terminal:", 
        "python test_delegated_access.py",
    ])

    # --- Lakehouses ---
    doc.add_heading("Available Lakehouses", level=2)

    _lakehouse(doc, "uzima_db_backup (primary)", "67596566-8ea9-4fd6-a451-ca9654aa4f10", [
        ("agents", "Study agent assignments"),
        ("dimdate", "Date dimension"),
        ("dimenrolledparticipants", "Enrolled participant details"),
        ("dimsleepdetailslogs", "Sleep log data"),
        ("dimsurveydictionary", "Survey question dictionary"),
        ("dimsurveyquestionresult", "Survey question results"),
        ("dimsurveyresults", "Survey responses"),
        ("dimsurveystepresults", "Survey step results"),
        ("dimsurveytask", "Survey task assignments"),
        ("factfitbitactivitieslogs", "Fitbit activity logs"),
        ("factfitbitdailydata", "Fitbit daily summaries"),
        ("factfitbitintraday", "Fitbit intraday data"),
        ("factfitbitintradaycombined", "Fitbit intraday (combined)"),
        ("factfitbitrestingheartrates", "Fitbit resting heart rate"),
        ("factfitbitsleeplogs", "Fitbit sleep logs"),
        ("qualtrics_hcw_student_survey", "Qualtrics survey data"),
        ("registeredparticipants", "Registered participants"),
        ("surveydata", "Survey data"),
        ("users_logonaudit", "User login audit"),
    ])

    _lakehouse(doc, "HCW_fitbit_data (shortcut)", "65058b40-a60c-4267-a882-9263e0ba0617", [
        ("fitbitactivitylogs", "Fitbit activity logs"),
        ("fitbitbodyweightlog", "Fitbit body weight logs"),
        ("fitbitdailydata", "Fitbit daily summaries"),
        ("fitbitdevices", "Fitbit device info"),
        ("fitbitfiles", "Fitbit raw files"),
    ])

    _lakehouse(doc, "Qualtrics (shortcut)", "8bb92d0b-3f94-4bd1-94d4-b31b088e9061", [
        ("dbo.aku_survey_responses_2026", "Qualtrics survey responses (256 columns)"),
    ])

    doc.add_paragraph(
        "\nTo access a specific lakehouse in Python:"
    )
    _code(doc, [
        'lh = FabricLakehouse(lakehouse_guid="65058b40-a60c-4267-a882-9263e0ba0617")',
    ])
    doc.add_paragraph("In R:")
    _code(doc, [
        'conn <- connect_to_fabric(',
        '  lakehouse_id = "65058b40-a60c-4267-a882-9263e0ba0617")',
    ])

    # --- Source ---
    doc.add_heading("Source Code", level=2)
    doc.add_paragraph("This setup is open-source. Code, docs, and runbooks:")
    _code(doc, ["https://github.com/AKU-CDIO/fabric-inbound-access"])

    # --- Contact ---
    doc.add_heading("Need help?", level=2)
    doc.add_paragraph("Contact Derick Imbati \u2014 derick.imbati@aku.edu")

    output_path = os.path.join(
        os.environ["TEMP"],
        "UZIMA_Instructions.docx",
    )
    doc.save(output_path)
    print(f"Saved: {output_path}")


def _lakehouse(doc, name, guid, tables):
    p = doc.add_paragraph()
    run = p.add_run(f"{name}")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(f"\n{guid}\n")
    for t, desc in tables:
        p2 = doc.add_paragraph(style="List Bullet")
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(1)
        run = p2.add_run(t + ": ")
        run.bold = True
        p2.add_run(desc)


def _code(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


if __name__ == "__main__":
    generate()
